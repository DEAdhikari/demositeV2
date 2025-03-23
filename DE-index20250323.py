import streamlit as st
import pandas as pd
from docx import Document
from fpdf import FPDF
import plotly.express as px
import matplotlib.pyplot as plt
from docx.shared import Inches
from datetime import datetime

#import matplotlib.pyplot as plt

st.set_page_config(page_title="Transaction Analyzer",
                   page_icon=":bar_chart:",
                   )

def generate_report(filtered_df1,graph_path):
    #print("Hello!")

    #st.title("Save Report as Word or PDF")

    
        # Code to generate Word report

    
         # Create a Word document
        doc = Document()
    
        # Add a title
        doc.add_heading("DataFrame Content", level=1)

        #st.dataframe(filtered_df1)
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        filename = f"PerformanceTestReport_{timestamp}.docx"
        # Add a table
        table = doc.add_table(rows=1, cols=len(filtered_df1.columns))
        table.style = "Table Grid"

        #print(filtered_df1.columns)  # Check available columns
        #print(filtered_df1.index)    # Check available row indices
    
        #if not filtered_df.empty:
        #table = doc.add_table(rows=1, cols=len(filtered_df.columns))
        
        hdr_cells = table.rows[0].cells
        for j, col in enumerate(filtered_df1.columns):
                hdr_cells[j].text = col
            
        for i, row in filtered_df1.iterrows():
                row_cells = table.add_row().cells
                for j, value in enumerate(row):
                    row_cells[j].text = str(value)
        
        # Add column headers
        #hdr_cells = table.rows[0].cells
        #for i, column in enumerate(filtered_df1.columns):
            #hdr_cells[i].text = str(column)

        # Add DataFrame rows to the table
        #for _, row in filtered_df1.iterrows():
            #row_cells = table.add_row().cells
        #for i, value in enumerate(row):
            #row_cells[i].text = str(value)

           
    
    # Add data rows
        #for index, row in filtered_df1.iterrows():
            #data_cells = table.add_row().cells
        #for i, cell_value in enumerate(row):
            #data_cells[i].text = str(cell_value)

        doc.add_heading("Graph Exported from Table", level=1)
        doc.add_picture(graph_path, width=Inches(5))

        # Save the document
        #doc.save("output.docx")
        doc.save(filename)
        st.success(f"Word document saved successfully")
        #print("DataFrame saved in 'output.docx'.")

             
    # Dropdown to select file format
    #file_format = st.selectbox("Choose the file format:", ["Word (.docx)", "PDF (.pdf)"])

   

# Load Excel file
uploaded_file = st.file_uploader("Choose a file", type=["csv", "xlsx", "txt"])

if uploaded_file is not None:
    df = pd.read_excel(uploaded_file, index_col=None)

    # Streamlit app
    st.title("Data Query and Graph Generator")

    # Select columns
    #st.sidebar.header("Please Filter Here:")
    #columns = st.sidebar.multiselect("Select columns to include", df.columns.tolist(), default=df.columns.tolist())

    # Select logical row names based on values in a specific column
    #st.sidebar.header("Please Filter Here:")
    #logical_column = st.sidebar.selectbox("Select column to filter rows", df.columns.tolist())
    #row_names = df[logical_column].unique().tolist()

    #st.sidebar.header("Please Filter Here:")
    #selected_rows = st.sidebar.multiselect("Select rows based on logical names", row_names, default=row_names)

     #Filter dataframe based on logical row names
    #filtered_df = df[df[logical_column].isin(selected_rows)][columns]

    # Display filtered dataframe without index
    #st.write("Filtered Dataframe")
    #st.dataframe(filtered_df)
    with st.expander("View DataFrame"):
        st.dataframe(df)

    # Get the list of columns
    columns = df.columns.tolist()

    # Select column to filter
    #st.sidebar.header("Please Filter Here:")
    #column_to_filter = st.sidebar.selectbox("Select column to filter", columns)
    column_to_filter = st.selectbox("Select column to filter", columns)

    # Get unique values in the selected column
    unique_values = df[column_to_filter].unique()

    # Select value to filter by
    #st.sidebar.header("Please Filter Here:")
    #filter_value = st.sidebar.selectbox("Select value to filter by", unique_values)
    filter_value = st.selectbox("Select value to filter by", unique_values)

    # Filter the DataFrame
    filtered_df = df[df[column_to_filter] == filter_value]

    # Display the filtered DataFrame
    #st.write("### Filtered Data")
    #st.dataframe(filtered_df)
    with st.expander("View Filetred DataFrame"):
        st.dataframe(filtered_df)

    del filtered_df['Status']

    st.sidebar.header("Please Filter Here:")
    selected_columns = st.sidebar.multiselect("Select columns to include", filtered_df.columns.tolist(), default=filtered_df.columns.tolist())

    if selected_columns:
        
        filtered_df1 = filtered_df[selected_columns]
    #st.write("Filtered Table:")
    #st.dataframe(filtered_df1)
    st.title(":bar_chart:  Transaction Analysis")
    #st.bar_chart(filtered_df1.set_index('TransactionName'), stack=False) 
        #print("Hello")

    fig, ax = plt.subplots()
    filtered_df1.plot(ax=ax)
    plt.title("Generated Graph")
    st.pyplot(fig)

    # Save Graph
    graph_path = "graph.png"
    fig.savefig(graph_path)

    if st.button('Generate Report'):
      #generate_report(filtered_df1) 
        #st.bar_chart(filtered_df1.set_index('TransactionName'), stack=False) 
        
        generate_report(filtered_df1,graph_path) 
    #fig = px.bar(df, x='TransactionName',title='Bar Chart Example')
    #fig.show()
else:
    st.write("")







    
        
  

